"""Extract plain text from a master resume (.docx, .pdf, or .txt)."""

from __future__ import annotations

from pathlib import Path

# Scoring/tailoring need real text; shorter usually means PDF extraction failed
# (common on Windows with image-only or oddly-encoded PDFs).
MIN_RESUME_CHARS = 80

_resume_cache: dict[str, str] = {}


def resume_text_looks_valid(text: str) -> bool:
    return len((text or "").strip()) >= MIN_RESUME_CHARS


def extract_text(path: Path) -> str:
    if not path.exists():
        raise FileNotFoundError(
            f"Master resume not found at {path}. Drop your resume there or set "
            f"MASTER_RESUME_PATH."
        )
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _from_docx(path)
    if suffix == ".pdf":
        return _from_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    raise ValueError(f"Unsupported resume format: {suffix} (use .docx, .pdf, or .txt)")


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _from_pdf_pypdf(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path.resolve()))
        parts: list[str] = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n".join(parts).strip()
    except Exception:
        return ""


def _from_pdf_pdfminer(path: Path) -> str:
    from pdfminer.high_level import extract_text as pdf_extract

    return (pdf_extract(str(path.resolve())) or "").strip()


def _from_pdf(path: Path) -> str:
    """Try pypdf first (often better on modern PDFs), then pdfminer.six."""
    pypdf_text = _from_pdf_pypdf(path)
    if resume_text_looks_valid(pypdf_text):
        return pypdf_text
    miner_text = _from_pdf_pdfminer(path)
    if len(miner_text.strip()) > len(pypdf_text.strip()):
        return miner_text
    return miner_text or pypdf_text


def cached_resume_text(path_str: str) -> str:
    key = str(Path(path_str).resolve())
    if key in _resume_cache:
        return _resume_cache[key]
    text = extract_text(Path(path_str))
    if resume_text_looks_valid(text):
        _resume_cache[key] = text
    return text


def clear_resume_cache() -> None:
    _resume_cache.clear()
