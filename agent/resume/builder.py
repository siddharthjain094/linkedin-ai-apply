"""Render tailored resume + cover letter documents and (optionally) compile to PDF."""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from rich.console import Console

console = Console()

# Bound every external conversion so a stuck Word/LibreOffice instance (e.g. an
# unanswerable macOS automation-permission dialog) can never hang the run.
_CONVERT_TIMEOUT_S = 90


def _find_soffice() -> str | None:
    """Locate the LibreOffice CLI across platforms.

    On Windows LibreOffice is almost never on PATH, so we also probe the standard
    install locations. ``LIBREOFFICE_PATH`` overrides everything if set."""
    override = os.environ.get("LIBREOFFICE_PATH")
    if override and Path(override).exists():
        return override

    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found

    candidates: list[str] = []
    if sys.platform == "win32":
        roots = [
            os.environ.get("PROGRAMFILES", r"C:\Program Files"),
            os.environ.get("PROGRAMFILES(X86)", r"C:\Program Files (x86)"),
        ]
        candidates += [str(Path(r) / "LibreOffice" / "program" / "soffice.exe")
                       for r in roots if r]
    elif sys.platform == "darwin":
        candidates.append("/Applications/LibreOffice.app/Contents/MacOS/soffice")
    else:  # linux/other
        candidates += ["/usr/bin/soffice", "/usr/local/bin/soffice",
                       "/opt/libreoffice/program/soffice", "/snap/bin/libreoffice"]

    return next((c for c in candidates if Path(c).exists()), None)


def _slug(text: str, maxlen: int = 60) -> str:
    keep = [c.lower() if c.isalnum() else "-" for c in text]
    s = "".join(keep)
    while "--" in s:
        s = s.replace("--", "-")
    return s.strip("-")[:maxlen] or "job"


def build_resume_docx(
    name: str,
    contact_line: str,
    summary: str,
    bullets: list[str],
    skills: list[str],
    base_resume_text: str,
    out_path: Path,
) -> Path:
    """Render a clean tailored resume docx.

    The tailored summary/bullets/skills lead the document; the full base resume
    text is appended so no real experience is lost.
    """
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    h = doc.add_heading(name, level=0)
    h.alignment = 1  # center
    if contact_line:
        c = doc.add_paragraph(contact_line)
        c.alignment = 1

    if summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(summary)

    if skills:
        doc.add_heading("Key Skills", level=1)
        doc.add_paragraph(", ".join(skills))

    if bullets:
        doc.add_heading("Selected Highlights", level=1)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("Experience & Background", level=1)
    for line in base_resume_text.splitlines():
        if line.strip():
            doc.add_paragraph(line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def build_cover_letter_docx(name: str, body: str, out_path: Path) -> Path:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)
    for para in body.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph("")
    doc.add_paragraph(name)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def to_pdf(docx_path: Path) -> Path | None:
    """Convert a docx to pdf. Returns None (and warns) if no converter succeeds.

    Order: LibreOffice headless first (reliable, no GUI), then Word via docx2pdf.
    Both are bounded by a timeout so a hung converter can't stall the pipeline."""
    pdf_path = docx_path.with_suffix(".pdf")

    # 1) LibreOffice / soffice headless (works on macOS, Windows, and Linux).
    soffice = _find_soffice()
    if soffice:
        # Use a private profile dir so conversion still works while a LibreOffice
        # GUI instance is open (otherwise headless mode errors with "already
        # running"). Path.as_uri() yields the file:// URI LO expects on every OS.
        profile = Path(tempfile.gettempdir()) / "linkedin_ai_apply_lo_profile"
        profile.mkdir(parents=True, exist_ok=True)
        try:
            subprocess.run(
                [soffice, f"-env:UserInstallation={profile.as_uri()}",
                 "--headless", "--norestore", "--convert-to", "pdf", "--outdir",
                 str(docx_path.parent), str(docx_path)],
                check=True, capture_output=True, timeout=_CONVERT_TIMEOUT_S,
            )
            if pdf_path.exists():
                return pdf_path
        except Exception:
            pass

    # 2) docx2pdf (Word on macOS/Windows). Run in a subprocess so a stuck Word
    #    automation / permission dialog is killed by the timeout instead of
    #    blocking the worker thread forever (docx2pdf.convert() has no timeout and
    #    silently swallows AppleScript errors, so we can't rely on it in-process).
    try:
        import docx2pdf  # noqa: F401  -- availability check only
        subprocess.run(
            [sys.executable, "-c",
             "import sys; from docx2pdf import convert; convert(sys.argv[1], sys.argv[2])",
             str(docx_path), str(pdf_path)],
            check=True, capture_output=True, timeout=_CONVERT_TIMEOUT_S,
        )
        if pdf_path.exists():
            return pdf_path
    except Exception:
        pass

    install_hint = {
        "darwin": "brew install --cask libreoffice",
        "win32": "winget install TheDocumentFoundation.LibreOffice",
    }.get(sys.platform, "install the libreoffice package")
    console.print(
        f"[yellow]PDF conversion unavailable[/]: could not render {docx_path.name} "
        f"to PDF. Install LibreOffice ([cyan]{install_hint}[/]) or set "
        "[cyan]RESUME_OUTPUT_FORMAT=docx[/]. Uploading the .docx instead.")
    return None


def job_basename(company: str, title: str) -> str:
    return f"{_slug(company)}_{_slug(title)}"
